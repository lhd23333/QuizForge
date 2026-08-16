"""生成 QuizForge 的 Pandoc Word 参考模板。

此脚本只在开发期维护样式资源，运行时导出不依赖 python-docx。
"""

from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parent.parent
DEFAULT_OUTPUT = PROJECT_ROOT / "assets" / "word-reference.docx"
CONTENT_WIDTH_DXA = 9746


def _set_run_font(run, *, east_asia: str, latin: str = "Calibri",
                  size: float, bold: bool = False,
                  color: str = "000000") -> None:
    run.font.name = latin
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    fonts.set(qn("w:ascii"), latin)
    fonts.set(qn("w:hAnsi"), latin)
    fonts.set(qn("w:eastAsia"), east_asia)


def _configure_style(document: Document, name: str, *, east_asia: str,
                     size: float, bold: bool = False,
                     alignment=None, before: float = 0, after: float = 4,
                     line: float = 1.25, color: str = "000000"):
    styles = document.styles
    style = styles[name] if name in styles else styles.add_style(
        name, WD_STYLE_TYPE.PARAGRAPH)
    style.base_style = styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = RGBColor.from_string(color)
    fonts = style.element.get_or_add_rPr().get_or_add_rFonts()
    fonts.set(qn("w:ascii"), "Calibri")
    fonts.set(qn("w:hAnsi"), "Calibri")
    fonts.set(qn("w:eastAsia"), east_asia)
    paragraph = style.paragraph_format
    paragraph.space_before = Pt(before)
    paragraph.space_after = Pt(after)
    paragraph.line_spacing = line
    if alignment is not None:
        paragraph.alignment = alignment
    return style


def _set_cell_width(cell, width_dxa: int) -> None:
    properties = cell._tc.get_or_add_tcPr()
    width = properties.find(qn("w:tcW"))
    if width is None:
        width = OxmlElement("w:tcW")
        properties.append(width)
    width.set(qn("w:w"), str(width_dxa))
    width.set(qn("w:type"), "dxa")
    margins = properties.find(qn("w:tcMar"))
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for side, value in (("top", 40), ("start", 60), ("bottom", 40), ("end", 60)):
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_table_geometry(table, widths: list[int], *, borders: bool) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    properties = table._tbl.tblPr
    table_width = properties.find(qn("w:tblW"))
    if table_width is None:
        table_width = OxmlElement("w:tblW")
        properties.append(table_width)
    table_width.set(qn("w:w"), str(sum(widths)))
    table_width.set(qn("w:type"), "dxa")
    layout = properties.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        properties.append(layout)
    layout.set(qn("w:type"), "fixed")
    borders_node = properties.find(qn("w:tblBorders"))
    if borders_node is None:
        borders_node = OxmlElement("w:tblBorders")
        properties.append(borders_node)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders_node.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders_node.append(node)
        node.set(qn("w:val"), "single" if borders else "nil")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:color"), "B7B7B7")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            _set_cell_width(cell, widths[min(index, len(widths) - 1)])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _header_footer_table(container, markers: list[str]) -> None:
    paragraph = container.paragraphs[0]
    paragraph._element.getparent().remove(paragraph._element)
    table = container.add_table(rows=1, cols=3, width=Inches(6.77))
    _set_table_geometry(table, [3249, 3248, 3249], borders=False)
    alignments = (
        WD_ALIGN_PARAGRAPH.LEFT,
        WD_ALIGN_PARAGRAPH.CENTER,
        WD_ALIGN_PARAGRAPH.RIGHT,
    )
    for cell, marker, alignment in zip(table.rows[0].cells, markers, alignments):
        target = cell.paragraphs[0]
        target.alignment = alignment
        target.paragraph_format.space_after = Pt(0)
        run = target.add_run(marker)
        _set_run_font(run, east_asia="宋体", size=9, color="666666")


def build_reference(output: Path) -> Path:
    document = Document()
    section = document.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.2677)
    section.page_height = Inches(11.6929)
    section.top_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.header_distance = Inches(0.4)
    section.footer_distance = Inches(0.4)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal_fonts = normal.element.get_or_add_rPr().get_or_add_rFonts()
    normal_fonts.set(qn("w:ascii"), "Calibri")
    normal_fonts.set(qn("w:hAnsi"), "Calibri")
    normal_fonts.set(qn("w:eastAsia"), "宋体")
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.25

    _configure_style(document, "ExamTitle", east_asia="黑体", size=18,
                     bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                     after=8, line=1.0)
    _configure_style(document, "ExamSubtitle", east_asia="宋体", size=11,
                     alignment=WD_ALIGN_PARAGRAPH.CENTER, after=8, line=1.0)
    _configure_style(document, "ExamNotice", east_asia="黑体", size=10,
                     bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                     after=4, line=1.0)
    _configure_style(document, "ExamNotes", east_asia="宋体", size=9.5,
                     after=6, line=1.2, color="333333")
    _configure_style(document, "QuestionType", east_asia="黑体", size=14,
                     bold=True, before=8, after=6, line=1.0)
    _configure_style(document, "Question", east_asia="宋体", size=10.5,
                     after=5, line=1.25)
    _configure_style(document, "Solution", east_asia="宋体", size=10,
                     after=5, line=1.2, color="333333")
    marker = _configure_style(document, "QuizForgeMarker", east_asia="宋体",
                              size=1, after=0, line=1.0, color="FFFFFF")
    marker.element.get_or_add_rPr().append(OxmlElement("w:vanish"))
    for name, alignment in (("ImageLeft", WD_ALIGN_PARAGRAPH.LEFT),
                            ("ImageCenter", WD_ALIGN_PARAGRAPH.CENTER),
                            ("ImageRight", WD_ALIGN_PARAGRAPH.RIGHT)):
        _configure_style(document, name, east_asia="宋体", size=10.5,
                         alignment=alignment, after=4, line=1.0)

    list_number = document.styles["List Number"]
    list_number.base_style = normal
    list_number.paragraph_format.left_indent = Inches(0.375)
    list_number.paragraph_format.first_line_indent = Inches(-0.25)
    list_number.paragraph_format.space_after = Pt(5)
    list_number.paragraph_format.line_spacing = 1.25

    _header_footer_table(section.header, [
        "QF_HEADER_LEFT", "QF_HEADER_CENTER", "QF_HEADER_RIGHT",
    ])
    _header_footer_table(section.footer, [
        "QF_FOOTER_LEFT", "QF_FOOTER_CENTER", "QF_FOOTER_RIGHT",
    ])

    document.add_paragraph("QuizForge 可编辑试卷", style="ExamTitle")
    document.add_paragraph("Word 样式参考页", style="ExamSubtitle")
    document.add_paragraph("一、示例题型", style="QuestionType")
    question = document.add_paragraph(style="List Number")
    question.add_run("已知 x = 1，请写出可继续编辑的答案。")
    document.add_paragraph("答案与解析：此处展示解析样式。", style="Solution")
    table = document.add_table(rows=2, cols=3)
    table.cell(0, 0).text = "姓名"
    table.cell(0, 1).text = "班级"
    table.cell(0, 2).text = "学号"
    _set_table_geometry(table, [3249, 3248, 3249], borders=True)

    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)
    return output


def main() -> int:
    parser = argparse.ArgumentParser(description="生成 QuizForge Word 参考模板")
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT)
    args = parser.parse_args()
    output = build_reference(args.output.resolve())
    print(output)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
